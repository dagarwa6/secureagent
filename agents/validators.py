"""
Agent Output Validators — Sprint 4
Validates all agent outputs for correctness, completeness, and hallucination indicators.
Used as a quality gate before final report generation.
"""

import re
import logging
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

VALID_NIST_FUNCTIONS = {"Govern", "Identify", "Protect", "Detect", "Respond", "Recover"}
VALID_PRIORITIES = {"Critical", "High", "Medium", "Low"}
VALID_GAP_TYPES = {"Systemic Structural Weakness", "Isolated Control Failure"}
MINIMUM_FINDINGS = 15
MINIMUM_REPORT_WORDS = 3500


@dataclass
class ValidationResult:
    passed: bool = True
    warnings: list[str] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)

    def fail(self, message: str):
        self.passed = False
        self.errors.append(message)

    def warn(self, message: str):
        self.warnings.append(message)

    def summary(self) -> str:
        status = "✅ PASSED" if self.passed else "❌ FAILED"
        lines = [status]
        for e in self.errors:
            lines.append(f"  ERROR: {e}")
        for w in self.warnings:
            lines.append(f"  WARN:  {w}")
        return "\n".join(lines)


class OutputValidator:
    """Validates all SecureAgent outputs before report generation."""

    def validate_all(self, state: dict) -> ValidationResult:
        """Run all validation checks on the pipeline state."""
        result = ValidationResult()

        self._validate_ingestion(state, result)
        self._validate_assessment(state, result)
        self._validate_risk_register(state, result)
        self._validate_fair(state, result)

        if result.passed:
            logger.info(f"Validation passed with {len(result.warnings)} warnings")
        else:
            logger.error(f"Validation failed: {len(result.errors)} errors, {len(result.warnings)} warnings")

        return result

    def _validate_ingestion(self, state: dict, result: ValidationResult):
        """Validate Ingestion Agent outputs."""
        assets = state.get("asset_inventory") or []
        policies = state.get("policy_refs") or []
        vendors = state.get("vendor_risks") or []
        summary = state.get("ingestion_summary", "")

        if len(assets) < 5:
            result.warn(f"Low asset count: {len(assets)} (expected 15+). Corpus may be incomplete.")
        if len(policies) < 3:
            result.warn(f"Low policy count: {len(policies)} (expected 10+).")
        if len(vendors) < 1:
            result.warn("No vendor risks extracted.")
        if len(summary) < 100:
            result.warn("Ingestion summary is very short — may indicate LLM failure.")

        # Check asset structure
        for asset in assets[:5]:
            if not isinstance(asset, dict):
                result.fail("Asset inventory contains non-dict entries")
                break

    def _validate_assessment(self, state: dict, result: ValidationResult):
        """Validate NIST CSF 2.0 assessment outputs."""
        nist_scores = state.get("nist_scores") or []

        if len(nist_scores) < 6:
            result.fail(f"NIST CSF assessment incomplete: only {len(nist_scores)}/6 functions scored")
            return

        seen_functions = set()
        for score in nist_scores:
            func = score.get("function", "")
            func_score = score.get("score", None)
            justification = score.get("score_justification", "")

            if func in seen_functions:
                result.warn(f"Duplicate NIST function: {func}")
            seen_functions.add(func)

            if func not in VALID_NIST_FUNCTIONS:
                result.warn(f"Unknown NIST function: {func}")

            if func_score is None:
                result.fail(f"Missing score for {func}")
            elif not (1.0 <= float(func_score) <= 5.0):
                result.fail(f"Score out of range for {func}: {func_score} (must be 1.0–5.0)")

            if len(str(justification)) < 30:
                result.warn(f"Short justification for {func} — may be vague")

        overall = state.get("overall_maturity_score")
        if overall is None or not (1.0 <= float(overall) <= 5.0):
            result.fail(f"Invalid overall maturity score: {overall}")

    def _validate_risk_register(self, state: dict, result: ValidationResult):
        """Validate risk register (minimum 15 findings, correct structure)."""
        findings = state.get("risk_findings") or []

        if len(findings) < MINIMUM_FINDINGS:
            result.fail(f"Risk register has only {len(findings)} findings (minimum {MINIMUM_FINDINGS} required)")
            return

        seen_ids = set()
        for i, finding in enumerate(findings):
            fid = finding.get("id", "")
            if fid in seen_ids:
                result.fail(f"Duplicate risk ID: {fid}")
            seen_ids.add(fid)

            # Check required fields
            for required_field in ["asset", "threat_scenario", "likelihood", "impact", "risk_score", "nist_function"]:
                if not finding.get(required_field):
                    result.warn(f"Finding {fid}: missing field '{required_field}'")

            # Validate math: risk_score = likelihood * impact
            likelihood = finding.get("likelihood", 0)
            impact = finding.get("impact", 0)
            risk_score = finding.get("risk_score", 0)
            expected_score = int(likelihood) * int(impact)
            if int(risk_score) != expected_score:
                result.warn(f"Finding {fid}: risk_score {risk_score} != {likelihood}×{impact}={expected_score}")

            # Validate NIST function
            nist_func = finding.get("nist_function", "")
            if nist_func and nist_func not in VALID_NIST_FUNCTIONS:
                result.warn(f"Finding {fid}: invalid NIST function '{nist_func}'")

            # Validate priority
            priority = finding.get("priority", "")
            if priority and priority not in VALID_PRIORITIES:
                result.warn(f"Finding {fid}: invalid priority '{priority}'")

        # Check for systemic findings
        systemic_count = sum(1 for f in findings if "Systemic" in str(f.get("gap_type", "")))
        if systemic_count == 0:
            result.warn("No systemic structural weaknesses identified — rubric requires classification")

    def _validate_fair(self, state: dict, result: ValidationResult):
        """Validate FAIR-lite results."""
        fair_results = state.get("fair_results") or []

        if len(fair_results) < 3:
            result.warn(f"Only {len(fair_results)} FAIR scenarios (expected 5)")
            return

        for fair in fair_results:
            ale = fair.get("ale_usd", None)
            tef = fair.get("tef_per_year", None)
            lm = fair.get("loss_magnitude_usd", None)
            ce = fair.get("control_effectiveness", None)

            if ale is None or tef is None or lm is None or ce is None:
                result.warn(f"FAIR result incomplete: {fair.get('risk_name', 'Unknown')}")
                continue

            # Verify ALE formula: ALE = TEF × LM × (1 - CE)
            expected_ale = float(tef) * float(lm) * (1 - float(ce))
            if abs(float(ale) - expected_ale) > 1.0:  # Allow $1 rounding tolerance
                result.warn(f"FAIR ALE formula mismatch for '{fair.get('risk_name', '')}': "
                            f"stored={ale:.0f}, expected={expected_ale:.0f}")

    def validate_report(self, docx_path: str) -> ValidationResult:
        """Validate generated DOCX report for completeness."""
        result = ValidationResult()

        if not docx_path:
            result.fail("Report path is empty — report may not have been generated")
            return result

        import os
        if not os.path.exists(docx_path):
            result.fail(f"Report file not found: {docx_path}")
            return result

        try:
            from docx import Document
            doc = Document(docx_path)

            # Count words (paragraphs + table cells)
            full_text = " ".join(para.text for para in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += " " + cell.text
            word_count = len(full_text.split())

            if word_count < MINIMUM_REPORT_WORDS:
                result.fail(f"Report too short: {word_count} words (minimum {MINIMUM_REPORT_WORDS} for 15+ pages)")

            # Check for required sections
            required_sections = [
                "Executive Summary",
                "Threat Model",
                "Risk Register",
                "Architecture",
                "Roadmap",
                "Governance",
            ]
            for section in required_sections:
                if section.lower() not in full_text.lower():
                    result.warn(f"Required section may be missing: '{section}'")

            # Check for placeholder text
            placeholders = re.findall(r"\[INSERT[^\]]*\]", full_text, re.IGNORECASE)
            if placeholders:
                result.warn(f"Placeholder text found: {placeholders[:3]}")

            file_size_kb = os.path.getsize(docx_path) / 1024
            if file_size_kb < 30:
                result.warn(f"Report file is very small ({file_size_kb:.0f} KB) — may be incomplete")

            result.warnings.append(f"Report validated: {word_count} words, {file_size_kb:.0f} KB")

        except Exception as e:
            result.fail(f"Could not open or validate report: {e}")

        return result


# ── QA Checklist ──────────────────────────────────────────────────────────────

RUBRIC_CHECKLIST = {
    "Sprint 1 — Project Charter (10 pts)": [
        "Engagement scope defined (in-scope + out-of-scope assets listed)",
        "Stakeholders identified with roles",
        "Milestones with dates (4 sprint gates)",
        "Success metrics defined (<10 min, 15+ findings, 15-20 pg report)",
        "Business objectives mapped to security risk",
    ],
    "Sprint 2 — Threat Modeling (20 pts)": [
        "STRIDE analysis complete (20+ threats across multiple assets)",
        "MITRE ATT&CK techniques mapped (15+ techniques)",
        "Cyber kill chain diagrammed (8 stages, healthcare ransomware scenario)",
        "Risk register has 15+ findings with all fields populated",
        "NIST CSF 2.0 scores for all 6 functions with justifications",
        "Systemic vs. isolated gap classification applied",
        "Healthcare industry benchmark comparison included",
    ],
    "Sprint 3 — Architecture & Governance (22 pts)": [
        "Target-state architecture designed (Zero Trust, IAM, EDR, SIEM)",
        "RACI matrix complete with all security functions",
        "18-month roadmap phased (Short/Mid/Long term)",
        "Budget estimates per phase ($755K total)",
        "FAIR-lite ALE for top 5 risks",
        "Change management plan (stakeholder comms, quick wins, resistance)",
        "Streamlit demo runs end-to-end",
        "DOCX report generates automatically",
    ],
    "Sprint 4 — Final Report + Defense (43 pts)": [
        "Written report 15-20+ pages",
        "Executive summary board-ready (no jargon)",
        "Visual elements: maturity radar, risk heat map, roadmap timeline",
        "15-20 slide board deck with strategic narrative",
        "2+ mock defenses rehearsed",
        "All agent decisions defensible without notes",
        "Report validated: 3500+ words, all sections present",
    ],
}


def print_qa_checklist():
    """Print the QA checklist to console for manual review."""
    print("\n" + "="*60)
    print("  SecureAgent — QA Rubric Checklist")
    print("="*60)
    for sprint, items in RUBRIC_CHECKLIST.items():
        print(f"\n{sprint}")
        for item in items:
            print(f"  ☐ {item}")
    print("\n" + "="*60 + "\n")


if __name__ == "__main__":
    print_qa_checklist()
