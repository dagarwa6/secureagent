"""
Report Generator — python-docx
Builds a professional consulting-grade DOCX report from AgentState data.

Report Structure (15–20 pages):
  1. Title Page
  2. Executive Summary
  3. Engagement Scope & Methodology
  4. Current-State Assessment (NIST CSF 2.0)
  5. Threat Model & Adversary Analysis
  6. Risk Register
  7. Target-State Architecture
  8. Governance & Policy Framework
  9. Implementation Roadmap
  10. FAIR-lite Risk Quantification
  11. Change Management
  12. Appendices
"""

import os
import logging
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# Brand colors (GSU blue-ish palette)
COLOR_PRIMARY = RGBColor(0x1A, 0x3A, 0x5C)   # Dark navy
COLOR_ACCENT = RGBColor(0xC8, 0x10, 0x26)     # Red accent
COLOR_HEADER_BG = RGBColor(0x1A, 0x3A, 0x5C)
COLOR_ALT_ROW = RGBColor(0xF0, 0xF4, 0xF8)    # Light blue-grey
COLOR_CRITICAL = RGBColor(0xC0, 0x00, 0x00)   # Red
COLOR_HIGH = RGBColor(0xFF, 0x6B, 0x00)        # Orange
COLOR_MEDIUM = RGBColor(0xFF, 0xC0, 0x00)      # Yellow
COLOR_LOW = RGBColor(0x00, 0x80, 0x00)         # Green


class ReportGenerator:
    def __init__(self):
        self.doc = Document()
        self._setup_document()

    def _setup_document(self):
        """Configure document margins and default styles."""
        from docx.shared import Inches
        section = self.doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def build_report(self, state: dict) -> str:
        """
        Build the complete consulting report from agent state.
        Returns path to saved .docx file.
        """
        logger.info("Building DOCX report...")

        self._add_title_page(state)
        self._add_page_break()
        self._add_table_of_contents()
        self._add_page_break()
        self._add_executive_summary(state)
        self._add_page_break()
        self._add_scope_and_methodology(state)
        self._add_scope_limitations(state)
        self._add_page_break()
        self._add_current_state_assessment(state)
        self._add_page_break()
        self._add_threat_model(state)
        self._add_page_break()
        self._add_risk_register(state)
        self._add_page_break()
        self._add_target_architecture(state)
        self._add_page_break()
        self._add_governance_section(state)
        self._add_page_break()
        self._add_roadmap_section(state)
        self._add_page_break()
        self._add_fair_quantification(state)
        self._add_page_break()
        self._add_change_management(state)
        self._add_page_break()
        self._add_appendices(state)
        self._add_page_break()
        self._add_data_provenance(state)

        # Save — use absolute path based on project root to avoid CWD issues
        _project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        output_dir = os.path.join(_project_root, "output")
        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(output_dir, f"MedBridge_SecurityAssessment_{timestamp}.docx")
        self.doc.save(output_path)
        logger.info(f"Report saved to {output_path}")
        return output_path

    # ── Section Builders ────────────────────────────────────────────────────────

    def _add_title_page(self, state: dict):
        """Add professional title page."""
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("MedBridge Health Systems")
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY

        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = subtitle.add_run("Cybersecurity Posture Assessment Report")
        run2.font.size = Pt(20)
        run2.font.color.rgb = COLOR_ACCENT

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        info_lines = [
            ("Prepared by:", "SecureAgent — AI-Powered Security Assessment Platform"),
            ("Engagement:", "CIS 8397 Capstone | Georgia State University"),
            ("Assessment Date:", datetime.now().strftime("%B %Y")),
            ("Classification:", "CONFIDENTIAL — For Executive Review Only"),
            ("NIST CSF Version:", "2.0"),
        ]
        for label, value in info_lines:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_label = p.add_run(f"{label} ")
            run_label.font.bold = True
            run_label.font.size = Pt(11)
            run_val = p.add_run(value)
            run_val.font.size = Pt(11)

    def _add_table_of_contents(self):
        self._add_heading("Table of Contents", level=1)
        toc_items = [
            "1. Executive Summary",
            "2. Engagement Scope & Methodology",
            "3. Current-State Security Assessment (NIST CSF 2.0)",
            "4. Threat Model & Adversary Analysis",
            "5. Risk Register (16 Findings)",
            "6. Target-State Architecture",
            "7. Governance & Policy Framework",
            "8. Implementation Roadmap (18 Months)",
            "9. FAIR-lite Risk Quantification",
            "10. Change Management",
            "Appendices",
        ]
        for item in toc_items:
            self.doc.add_paragraph(item, style="List Number")

    def _add_executive_summary(self, state: dict):
        self._add_heading("1. Executive Summary", level=1)

        exec_summary = state.get("executive_summary", "")
        if exec_summary:
            self.doc.add_paragraph(exec_summary)
        else:
            self.doc.add_paragraph(
                "MedBridge Health Systems presents a high-risk cybersecurity posture with an overall NIST CSF 2.0 maturity "
                "score of 1.88/5.0, significantly below the healthcare industry benchmark of 2.1/5.0. The organization "
                "faces critical exposure to ransomware ($1.07M Annual Loss Expectancy), PHI data breach ($720K ALE), and "
                "supply chain compromise risks — driven by foundational gaps including no CISO, no SIEM, no EDR, and MFA "
                "coverage at only 35%.\n\n"
                "The three most critical findings are: (1) Absence of security leadership and governance — no CISO position "
                "exists, leaving security decisions ad-hoc and underfunded; (2) Universal MFA gap — clinical staff at 12% "
                "enrollment creates a high-probability ransomware entry point as evidenced by the 2023 credential phishing "
                "incident; (3) No detection capability — the absence of a SIEM and behavioral EDR results in mean time to "
                "detect of 6+ hours to 11 days, enabling extended attacker dwell time for data exfiltration and ransomware staging.\n\n"
                "SecureAgent recommends an 18-month, phased security transformation investment of $755,000, which reduces "
                "the organization's top-5 risk Annual Loss Expectancy from $2.46M to approximately $200K — a 3:1 return "
                "on investment within the first 18 months, while bringing MedBridge into alignment with HIPAA Security Rule "
                "requirements and NIST CSF 2.0 healthcare benchmarks."
            )

        # Key metrics callout table
        self._add_heading("Key Assessment Metrics", level=2)
        table = self.doc.add_table(rows=5, cols=2)
        table.style = "Table Grid"
        metrics = [
            ("Overall NIST CSF 2.0 Score", f"{state.get('overall_maturity_score', 1.88):.2f}/5.0 (Healthcare Benchmark: 2.1)"),
            ("Total Risk Findings", f"{len(state.get('risk_findings', []) or [])} findings (5 Critical, 6 High, 4 Medium, 1 Low)"),
            ("Top Annual Loss Exposure", "$2.46M across top 5 risk scenarios (FAIR-lite)"),
            ("Recommended 18-Month Investment", "$755,000 (ROI: $2.26M ALE reduction)"),
            ("Critical Gap Summary", "No CISO | No SIEM | No EDR | MFA 35% | EOL Firewall"),
        ]
        for i, (label, value) in enumerate(metrics):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].text = value

    def _add_scope_and_methodology(self, state: dict):
        self._add_heading("2. Engagement Scope & Methodology", level=1)

        self._add_heading("2.1 Client Profile", level=2)
        self.doc.add_paragraph(
            "MedBridge Health Systems is a mid-size healthcare organization operating 6 facilities across the Atlanta "
            "metropolitan area, serving approximately 340,000 patients. The organization employs 1,200 staff including "
            "680 clinical personnel, 420 administrative staff, and 100 IT and support personnel. MedBridge operates under "
            "HIPAA and HITECH regulatory frameworks and relies on Epic EHR as its core clinical system in a hybrid "
            "Azure + on-premises environment."
        )

        self._add_heading("2.2 Scope", level=2)
        scope_items = [
            "Network infrastructure (perimeter, internal, Azure cloud)",
            "Epic EHR application and database infrastructure",
            "Active Directory and Azure AD identity environment",
            "40 Windows servers, 12 Linux servers, 492 managed endpoints",
            "Three third-party vendor connections (LabConnect, RadCloud, PaySync)",
            "22 security policies (including 8 identified as missing)",
            "Security governance structure and organizational roles",
        ]
        for item in scope_items:
            self.doc.add_paragraph(item, style="List Bullet")

        self._add_heading("2.3 Methodology", level=2)
        self.doc.add_paragraph(
            "SecureAgent employs a five-phase AI-assisted assessment pipeline: (1) Document Ingestion — structured "
            "extraction of assets, policies, and vendor relationships from organizational documents; (2) Threat Modeling "
            "— STRIDE analysis and MITRE ATT&CK technique mapping; (3) Current-State Assessment — NIST CSF 2.0 maturity "
            "scoring across all 6 functions with CIS Controls v8 cross-reference; (4) Gap Analysis — control gap "
            "identification with FAIR-lite financial quantification; (5) Report Generation — synthesis of all findings "
            "into this board-ready consulting report."
        )

    def _add_scope_limitations(self, state: dict):
        """Add Scope & Limitations section for intellectual honesty."""
        self._add_heading("2.4 Scope & Limitations", level=2)
        self.doc.add_paragraph(
            "This assessment is subject to the following limitations, which should be considered "
            "when interpreting findings and recommendations:"
        )
        limitations = [
            "Document-Based Analysis Only: This assessment is based entirely on analysis of provided "
            "organizational documents. No live infrastructure scanning, vulnerability testing, penetration "
            "testing, staff interviews, or physical facility inspection was performed.",
            "Corpus-Limited Findings: Findings are limited to what is documented in the provided corpus. "
            "Undocumented risks, shadow IT, and controls not reflected in the documents are not captured.",
            "FAIR Risk Estimates: Financial risk estimates use Monte Carlo simulation with triangular "
            "distributions based on industry benchmarks and documented evidence. These should not replace "
            "actuarial analysis or formal FAIR assessments with subject-matter-expert-calibrated inputs.",
            "Maturity Scores: NIST CSF 2.0 maturity scores reflect document-based evidence only and may "
            "differ from scores derived through hands-on technical testing, control validation, or staff interviews.",
            "AI-Assisted Analysis: This assessment was produced using an AI-assisted pipeline (LLM-driven). "
            "While the system is designed to reduce human bias, AI-generated findings should be reviewed by "
            "qualified cybersecurity professionals before informing investment decisions.",
            "Point-in-Time Assessment: This report reflects the organization's documented posture at the time "
            "of analysis. The threat landscape and organizational controls evolve continuously.",
            "Supplementary Tool: SecureAgent is designed to supplement, not replace, human-led security "
            "assessments. It accelerates the initial assessment process and provides a structured framework "
            "for further investigation.",
        ]
        for limitation in limitations:
            self.doc.add_paragraph(limitation, style="List Bullet")

    def _add_current_state_assessment(self, state: dict):
        self._add_heading("3. Current-State Security Assessment", level=1)
        self.doc.add_paragraph(
            "SecureAgent assessed MedBridge's cybersecurity posture against the NIST Cybersecurity Framework 2.0, "
            f"assigning maturity scores on a 1–5 scale. The overall weighted score of "
            f"{state.get('overall_maturity_score', 1.88):.2f}/5.0 is below the healthcare industry benchmark "
            f"of {state.get('industry_benchmark', 2.1)}/5.0 (CISA 2024)."
        )

        # NIST Scores Table
        self._add_heading("3.1 NIST CSF 2.0 Maturity Scores", level=2)
        nist_scores = state.get("nist_scores") or []
        if nist_scores:
            table = self.doc.add_table(rows=len(nist_scores) + 1, cols=4)
            table.style = "Table Grid"
            headers = ["Function", "Score", "Maturity Level", "Critical Gap"]
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                self._set_cell_bg(cell, COLOR_HEADER_BG)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            for i, score in enumerate(nist_scores):
                row = table.rows[i + 1]
                row.cells[0].text = f"{score.get('function', 'Unknown')} ({score.get('function_id', '')})"
                row.cells[1].text = f"{score.get('score', '?'):.1f}/5.0"
                row.cells[2].text = score.get("maturity_level", "Unknown")
                gaps = score.get("key_gaps", [])
                row.cells[3].text = gaps[0] if gaps else "See full assessment"

    def _add_threat_model(self, state: dict):
        self._add_heading("4. Threat Model & Adversary Analysis", level=1)

        self._add_heading("4.1 Threat Actor Profile", level=2)
        self.doc.add_paragraph(
            "Healthcare is the #1 ransomware target sector (HC3 2024). MedBridge faces threats primarily from "
            "ransomware affiliates (FIN12/ALPHV) specializing in healthcare targets, malicious insiders with "
            "privileged EHR access, and opportunistic actors exploiting unpatched systems."
        )
        actors = state.get("top_threat_actors") or []
        for actor in actors:
            self.doc.add_paragraph(actor, style="List Bullet")

        self._add_heading("4.2 Cyber Kill Chain", level=2)
        self.doc.add_paragraph(
            "The following ransomware kill chain represents the most probable attack scenario for MedBridge, "
            "based on observed healthcare threat actor TTPs and MedBridge's specific vulnerabilities:"
        )
        kill_chain = state.get("kill_chain") or []
        for stage in kill_chain:
            self._add_heading(stage.get("stage", "Stage"), level=3)
            self.doc.add_paragraph(stage.get("description", ""))
            mapped = stage.get("mapped_techniques", [])
            if mapped:
                p = self.doc.add_paragraph("MITRE ATT&CK: ")
                p.add_run(", ".join(mapped)).font.italic = True

        self._add_heading("4.3 STRIDE Threat Summary", level=2)
        stride_threats = state.get("stride_threats") or []
        if stride_threats:
            # Show top 10 STRIDE threats
            top_stride = sorted(stride_threats,
                                key=lambda x: x.get("likelihood", 0) * x.get("impact", 0),
                                reverse=True)[:10]
            table = self.doc.add_table(rows=len(top_stride) + 1, cols=4)
            table.style = "Table Grid"
            headers = ["Asset", "STRIDE Category", "Threat", "Likelihood"]
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                self._set_cell_bg(cell, COLOR_HEADER_BG)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for i, threat in enumerate(top_stride):
                row = table.rows[i + 1]
                row.cells[0].text = str(threat.get("asset", ""))[:40]
                row.cells[1].text = str(threat.get("stride_category", ""))
                row.cells[2].text = str(threat.get("threat_description", ""))[:80]
                row.cells[3].text = str(threat.get("likelihood", "?"))
        else:
            self.doc.add_paragraph("Threat model data not available.")

    def _add_risk_register(self, state: dict):
        self._add_heading("5. Risk Register", level=1)
        findings = state.get("risk_findings") or []
        systemic = state.get("systemic_count", 0) or 0
        isolated = state.get("isolated_count", 0) or 0

        self.doc.add_paragraph(
            f"SecureAgent identified {len(findings)} prioritized risk findings for MedBridge Health Systems. "
            f"Of these, {systemic} represent Systemic Structural Weaknesses (cross-domain governance failures) "
            f"and {isolated} represent Isolated Control Failures (tactical gaps)."
        )

        if findings:
            table = self.doc.add_table(rows=len(findings) + 1, cols=6)
            table.style = "Table Grid"
            headers = ["ID", "Asset", "Threat Scenario", "Risk Score", "Gap Type", "Priority"]
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                self._set_cell_bg(cell, COLOR_HEADER_BG)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            for i, finding in enumerate(findings[:20]):  # Max 20 in table
                row = table.rows[i + 1]
                row.cells[0].text = str(finding.get("id", ""))
                row.cells[1].text = str(finding.get("asset", ""))[:30]
                row.cells[2].text = str(finding.get("threat_scenario", ""))[:80]
                row.cells[3].text = str(finding.get("risk_score", ""))
                gap_type = str(finding.get("gap_type", ""))
                row.cells[4].text = "Systemic" if "Systemic" in gap_type else "Isolated"
                row.cells[5].text = str(finding.get("priority", ""))

    def _add_target_architecture(self, state: dict):
        self._add_heading("6. Target-State Security Architecture", level=1)
        arch = state.get("architecture_recommendations", "")
        if arch:
            # Parse markdown and add as paragraphs
            for line in arch.split("\n"):
                line = line.strip()
                if line.startswith("## "):
                    self._add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    self._add_heading(line[4:], level=3)
                elif line.startswith("**") and line.endswith("**"):
                    p = self.doc.add_paragraph()
                    p.add_run(line[2:-2]).font.bold = True
                elif line.startswith("- "):
                    self.doc.add_paragraph(line[2:], style="List Bullet")
                elif line:
                    self.doc.add_paragraph(line)

    def _add_governance_section(self, state: dict):
        self._add_heading("7. Governance & Policy Framework", level=1)
        gov = state.get("governance_output") or {}

        self._add_heading("7.1 Recommended Governance Structure", level=2)
        gov_structure = gov.get("governance_structure", "")
        if gov_structure:
            self.doc.add_paragraph(gov_structure)

        # RACI table
        self._add_heading("7.2 RACI Matrix — Security Accountabilities", level=2)
        raci = gov.get("raci_matrix", [])
        if raci:
            table = self.doc.add_table(rows=len(raci) + 1, cols=5)
            table.style = "Table Grid"
            headers = ["Security Function", "Responsible", "Accountable", "Consulted", "Informed"]
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                self._set_cell_bg(cell, COLOR_HEADER_BG)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for i, entry in enumerate(raci):
                row = table.rows[i + 1]
                row.cells[0].text = str(entry.get("security_function", ""))
                row.cells[1].text = str(entry.get("responsible", ""))
                row.cells[2].text = str(entry.get("accountable", ""))
                row.cells[3].text = str(entry.get("consulted", ""))
                row.cells[4].text = str(entry.get("informed", ""))

        # Priority policy updates
        self._add_heading("7.3 Priority Policy Actions", level=2)
        policy_updates = gov.get("policy_updates", [])
        if policy_updates:
            table = self.doc.add_table(rows=len(policy_updates) + 1, cols=4)
            table.style = "Table Grid"
            headers = ["Policy", "Status", "Priority", "Target Timeline"]
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                self._set_cell_bg(cell, COLOR_HEADER_BG)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for i, policy in enumerate(policy_updates):
                row = table.rows[i + 1]
                row.cells[0].text = str(policy.get("policy", ""))
                row.cells[1].text = str(policy.get("status", ""))
                row.cells[2].text = str(policy.get("priority", ""))
                row.cells[3].text = str(policy.get("timeline", ""))

        # Training plan
        self._add_heading("7.4 Security Training & Awareness Plan", level=2)
        training_plan = gov.get("training_plan", [])
        if training_plan:
            self.doc.add_paragraph(
                "MedBridge requires a comprehensive, role-based security training program to address the "
                "34% phishing click rate and 61% training completion rate identified during the assessment. "
                "The following training plan covers all workforce segments with appropriate frequency and format."
            )
            table = self.doc.add_table(rows=len(training_plan) + 1, cols=4)
            table.style = "Table Grid"
            headers = ["Audience", "Training", "Frequency", "Format"]
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                self._set_cell_bg(cell, COLOR_HEADER_BG)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for i, item in enumerate(training_plan):
                row = table.rows[i + 1]
                row.cells[0].text = str(item.get("audience", ""))
                row.cells[1].text = str(item.get("training", ""))
                row.cells[2].text = str(item.get("frequency", ""))
                row.cells[3].text = str(item.get("format", ""))

        # Executive KPIs
        self._add_heading("7.5 Executive Security KPIs", level=2)
        kpis = gov.get("kpis", [])
        if kpis:
            self.doc.add_paragraph(
                "The following Key Performance Indicators provide the CISO and Board with measurable "
                "progress tracking against the security transformation roadmap. Each KPI includes current "
                "baseline, target state, and reporting cadence for executive dashboards."
            )
            table = self.doc.add_table(rows=len(kpis) + 1, cols=5)
            table.style = "Table Grid"
            headers = ["Metric", "Current", "Target", "Timeline", "Reporting"]
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                self._set_cell_bg(cell, COLOR_HEADER_BG)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for i, kpi in enumerate(kpis):
                row = table.rows[i + 1]
                row.cells[0].text = str(kpi.get("metric", ""))
                row.cells[1].text = str(kpi.get("current", ""))
                row.cells[2].text = str(kpi.get("target", ""))
                row.cells[3].text = str(kpi.get("timeline", ""))
                row.cells[4].text = str(kpi.get("reporting_freq", ""))

    def _add_roadmap_section(self, state: dict):
        self._add_heading("8. Implementation Roadmap (18 Months)", level=1)
        roadmap = state.get("roadmap") or []

        total_budget = sum(p.get("budget_estimate_usd", 0) for p in roadmap)
        self.doc.add_paragraph(
            f"The recommended 18-month implementation roadmap addresses all Critical and High risk findings in "
            f"a phased approach, with a total estimated investment of ${total_budget:,}. Phases are designed "
            f"to deliver measurable risk reduction at each gate, with the highest-risk items addressed first."
        )

        for phase in roadmap:
            self._add_heading(f"{phase.get('phase', 'Phase')}", level=2)
            self.doc.add_paragraph(
                f"Timeframe: {phase.get('timeframe', '')} | "
                f"Budget: ${phase.get('budget_estimate_usd', 0):,} | "
                f"Theme: {phase.get('theme', '')}"
            )

            initiatives = phase.get("initiatives", [])
            if initiatives:
                table = self.doc.add_table(rows=len(initiatives) + 1, cols=4)
                table.style = "Table Grid"
                headers = ["Initiative", "Timeline", "Cost (USD)", "Priority"]
                for j, h in enumerate(headers):
                    cell = table.rows[0].cells[j]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    self._set_cell_bg(cell, COLOR_HEADER_BG)
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                for i, initiative in enumerate(initiatives):
                    row = table.rows[i + 1]
                    row.cells[0].text = str(initiative.get("name", ""))
                    row.cells[1].text = str(initiative.get("timeline", ""))
                    row.cells[2].text = f"${initiative.get('cost_usd', 0):,}"
                    row.cells[3].text = str(initiative.get("priority", ""))
                self.doc.add_paragraph()

            if phase.get("success_criteria"):
                self.doc.add_paragraph("Phase Success Criteria:")
                for criterion in phase["success_criteria"]:
                    self.doc.add_paragraph(criterion, style="List Bullet")
            self.doc.add_paragraph()

        # Vendor-grounded budget breakdown
        self._add_heading("8.4 Budget Breakdown with Vendor References", level=2)
        self.doc.add_paragraph(
            "The following table provides vendor-level cost detail for each major line item. "
            "Estimates are vendor-neutral but reference market pricing tiers to ensure defensibility. "
            "All costs are annualized or amortized over the 18-month roadmap period."
        )
        budget_items = [
            ("CISO Hire (Full-Time)", "$180K-$220K/yr", "Salary.com healthcare CISO median: $195K base + benefits. Alternative: vCISO at $15-25K/month."),
            ("EDR (492 Endpoints)", "$42K-$54K/yr", "Microsoft Defender for Endpoint P2: $5.20/user/mo (M365 E5) or CrowdStrike Falcon Go: ~$8.99/endpoint/mo."),
            ("SIEM (Azure Sentinel)", "$96K-$144K/yr", "Azure Sentinel: ~$2.46/GB ingested. Est. 5-10 GB/day = $4.5-9K/mo. Includes log analytics workspace."),
            ("MFA Rollout (1,200 users)", "$12K-$18K/yr", "Azure AD P2: $9/user/mo (incl. in E5). Duo MFA: $3-9/user/mo. FIDO2 keys: $25-50/key for privileged users."),
            ("Security Awareness Training", "$18K-$36K/yr", "KnowBe4: $15-30/user/yr for 1,200 users. Proofpoint SA: similar tier. Includes phishing simulations."),
            ("NGFW Replacement (EOL ASA)", "$60K-$80K", "Palo Alto PA-400: $25-40K + licensing. Fortinet 100F: $15-25K + FortiGuard bundle. One-time CapEx."),
            ("IR Retainer + IRP Update", "$25K-$40K/yr", "CrowdStrike/Mandiant IR retainer: $25-40K/yr for 40-80 pre-purchased hours."),
            ("BCP/DR Testing", "$30K-$50K", "Professional services for Epic recovery validation, BCP tabletop, and DRP development."),
            ("Vuln. Mgmt Expansion", "$30K-$45K/yr", "Tenable.io: ~$2,275/asset/yr. Qualys VMDR: similar. Expand to Linux and cloud workloads."),
            ("Governance & Policy", "$30K-$40K", "External GRC consultant for 8 missing policies + HIPAA risk analysis. 2-3 month engagement."),
        ]
        budget_table = self.doc.add_table(rows=len(budget_items) + 1, cols=3)
        budget_table.style = "Table Grid"
        for j, h in enumerate(["Line Item", "Estimated Cost", "Vendor Reference / Basis"]):
            cell = budget_table.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True
            self._set_cell_bg(cell, COLOR_HEADER_BG)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for i, (item, cost, ref) in enumerate(budget_items):
            row = budget_table.rows[i + 1]
            row.cells[0].text = item
            row.cells[1].text = cost
            row.cells[2].text = ref

        self._add_heading("8.5 Budget Assumptions", level=2)
        assumptions = [
            "Pricing reflects 2025-2026 list rates; enterprise negotiation or M365 E5 bundling (includes MDE, Azure AD P2, Sentinel credits) may reduce total cost.",
            "CISO cost assumes full-time hire in Atlanta metro market; vCISO is a lower-cost interim option.",
            "SIEM cost is consumption-based (pay-per-GB) and varies with log volume and retention.",
            "Hardware (NGFW, FIDO2 keys) are one-time CapEx; software costs are annualized subscriptions.",
            "Professional services estimated from typical healthcare engagement rates.",
            "Internal IT staff time for implementation excluded — absorbed by existing 7 FTE IT team with CISO oversight.",
        ]
        for assumption in assumptions:
            self.doc.add_paragraph(assumption, style="List Bullet")

    def _add_fair_quantification(self, state: dict):
        self._add_heading("9. FAIR Risk Quantification", level=1)
        self.doc.add_paragraph(
            "The following analysis presents Annual Loss Expectancy (ALE) estimates for the top 5 MedBridge risk "
            "scenarios using the FAIR (Factor Analysis of Information Risk) model. Point estimates use the formula "
            "ALE = TEF × LM × (1 - CE). Additionally, a Monte Carlo simulation (10,000 iterations with "
            "triangular distributions) provides confidence intervals — the 10th and 90th percentile values "
            "represent the optimistic and pessimistic bounds, capturing the inherent uncertainty in risk quantification."
        )

        fair_results = state.get("fair_results") or []
        if fair_results:
            # Point estimate table
            self._add_heading("9.1 Point Estimates", level=2)
            table = self.doc.add_table(rows=len(fair_results) + 2, cols=5)
            table.style = "Table Grid"
            headers = ["Risk Scenario", "TEF/yr", "Loss Magnitude", "Control Eff.", "ALE/yr"]
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                self._set_cell_bg(cell, COLOR_HEADER_BG)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            total_ale = 0
            for i, result in enumerate(fair_results):
                row = table.rows[i + 1]
                row.cells[0].text = str(result.get("risk_name", ""))[:50]
                row.cells[1].text = str(result.get("tef_per_year", ""))
                row.cells[2].text = f"${result.get('loss_magnitude_usd', 0):,.0f}"
                row.cells[3].text = f"{result.get('control_effectiveness', 0)*100:.0f}%"
                row.cells[4].text = result.get("ale_formatted", "")
                total_ale += result.get("ale_usd", 0)

            # Total row
            total_row = table.rows[-1]
            total_row.cells[0].text = "TOTAL TOP-5 ANNUAL LOSS EXPOSURE"
            total_row.cells[0].paragraphs[0].runs[0].font.bold = True
            total_row.cells[4].text = f"${total_ale:,.0f}"
            total_row.cells[4].paragraphs[0].runs[0].font.bold = True

            # Monte Carlo confidence interval table
            has_mc = any(r.get("ale_p10") is not None for r in fair_results)
            if has_mc:
                self._add_heading("9.2 Monte Carlo Confidence Intervals (10,000 iterations)", level=2)
                self.doc.add_paragraph(
                    "The table below shows the range of probable ALE values from Monte Carlo simulation. "
                    "The P10 value means there is a 90% chance ALE will exceed this amount; P90 means "
                    "there is only a 10% chance ALE will exceed this amount."
                )
                mc_table = self.doc.add_table(rows=len(fair_results) + 1, cols=4)
                mc_table.style = "Table Grid"
                mc_headers = ["Risk Scenario", "P10 (Optimistic)", "Median", "P90 (Pessimistic)"]
                for j, h in enumerate(mc_headers):
                    cell = mc_table.rows[0].cells[j]
                    cell.text = h
                    cell.paragraphs[0].runs[0].font.bold = True
                    self._set_cell_bg(cell, COLOR_HEADER_BG)
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                for i, result in enumerate(fair_results):
                    row = mc_table.rows[i + 1]
                    row.cells[0].text = str(result.get("risk_name", ""))[:50]
                    p10 = result.get("ale_p10")
                    median = result.get("ale_median")
                    p90 = result.get("ale_p90")
                    row.cells[1].text = f"${p10:,.0f}" if p10 is not None else "N/A"
                    row.cells[2].text = f"${median:,.0f}" if median is not None else "N/A"
                    row.cells[3].text = f"${p90:,.0f}" if p90 is not None else "N/A"

    def _add_change_management(self, state: dict):
        self._add_heading("10. Change Management", level=1)
        cm = state.get("change_management") or {}

        self._add_heading("10.1 90-Day Quick Wins (Zero or Near-Zero Cost)", level=2)
        quick_wins = cm.get("quick_wins_90_days", [])
        for win in quick_wins:
            self.doc.add_paragraph(win, style="List Bullet")

        self._add_heading("10.2 Resistance Mitigation", level=2)
        resistance = cm.get("resistance_mitigation", [])
        for item in resistance:
            self._add_heading(item.get("resistance", ""), level=3)
            self.doc.add_paragraph(item.get("mitigation", ""))

        self._add_heading("10.3 Stakeholder Communication Plan", level=2)
        comms = cm.get("stakeholder_communications", [])
        if comms:
            table = self.doc.add_table(rows=len(comms) + 1, cols=4)
            table.style = "Table Grid"
            headers = ["Audience", "Key Message", "Frequency", "Owner"]
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
                self._set_cell_bg(cell, COLOR_HEADER_BG)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for i, comm in enumerate(comms):
                row = table.rows[i + 1]
                row.cells[0].text = str(comm.get("audience", ""))
                row.cells[1].text = str(comm.get("message", ""))[:80]
                row.cells[2].text = str(comm.get("frequency", ""))
                row.cells[3].text = str(comm.get("owner", ""))

    def _add_appendices(self, state: dict):
        self._add_heading("Appendices", level=1)

        self._add_heading("Appendix A: Complete Risk Register", level=2)
        findings = state.get("risk_findings") or []
        self.doc.add_paragraph(f"Full risk register: {len(findings)} findings")
        for finding in findings:
            self._add_heading(f"{finding.get('id', '')} — {finding.get('threat_scenario', '')[:60]}", level=3)
            details = [
                f"Asset: {finding.get('asset', '')}",
                f"Likelihood: {finding.get('likelihood', '')}/5 | Impact: {finding.get('impact', '')}/5 | Risk Score: {finding.get('risk_score', '')}",
                f"Gap Type: {finding.get('gap_type', '')}",
                f"Control Gap: {finding.get('control_gap', '')}",
                f"Recommended Control: {finding.get('recommended_control', '')}",
                f"NIST Function: {finding.get('nist_function', '')} | Business Objective: {finding.get('business_objective_at_risk', '')}",
            ]
            for detail in details:
                self.doc.add_paragraph(detail, style="List Bullet")

        self._add_heading("Appendix B: CIS Controls v8 Mapping", level=2)
        cis_controls = state.get("cis_controls_mapped") or []
        if cis_controls:
            table = self.doc.add_table(rows=len(cis_controls) + 1, cols=4)
            table.style = "Table Grid"
            headers = ["CIS Control", "Title", "Status", "NIST Function"]
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
            for i, ctrl in enumerate(cis_controls):
                row = table.rows[i + 1]
                row.cells[0].text = str(ctrl.get("control_id", ""))
                row.cells[1].text = str(ctrl.get("title", ""))
                row.cells[2].text = str(ctrl.get("status", ""))
                row.cells[3].text = str(ctrl.get("nist_function", ""))

    def _add_data_provenance(self, state: dict):
        """Add Data Provenance appendix showing which sections used LLM vs fallback data."""
        self._add_heading("Appendix C: Data Provenance", level=2)
        self.doc.add_paragraph(
            "This appendix documents which sections of this report were generated using "
            "live LLM analysis versus static fallback data. When the LLM is unavailable or "
            "returns an error, the pipeline falls back to pre-defined templates. Sections "
            "marked as 'Fallback (Static)' should be reviewed with additional scrutiny, as "
            "they were not derived from AI analysis of the provided corpus documents."
        )
        fallback_flags = state.get("fallback_flags") or {}
        agent_labels = {
            "ingestion": "Agent 1: Document Ingestion",
            "threat_modeling": "Agent 2: Threat Modeling (STRIDE + MITRE)",
            "assessment": "Agent 3: NIST CSF 2.0 Assessment",
            "gap_analysis": "Agent 4: Gap Analysis & Risk Register",
            "report_generation": "Agent 5: Report Generation",
        }
        table = self.doc.add_table(rows=len(agent_labels) + 1, cols=3)
        table.style = "Table Grid"
        headers = ["Pipeline Agent", "Data Source", "Status"]
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True
            self._set_cell_bg(cell, COLOR_HEADER_BG)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for i, (key, label) in enumerate(agent_labels.items()):
            row = table.rows[i + 1]
            row.cells[0].text = label
            source = fallback_flags.get(key, "unknown")
            if source == "llm_generated":
                row.cells[1].text = "LLM-Generated"
                row.cells[2].text = "✓ Live AI Analysis"
            elif source == "fallback_static":
                row.cells[1].text = "Fallback (Static)"
                row.cells[2].text = "⚠ Pre-defined Templates"
            else:
                row.cells[1].text = "Not Executed"
                row.cells[2].text = "— Skipped"

    # ── Helpers ──────────────────────────────────────────────────────────────────

    def _add_heading(self, text: str, level: int = 1):
        heading = self.doc.add_heading(text, level=level)
        if level == 1:
            for run in heading.runs:
                run.font.color.rgb = COLOR_PRIMARY
        elif level == 2:
            for run in heading.runs:
                run.font.color.rgb = COLOR_ACCENT

    def _add_page_break(self):
        self.doc.add_page_break()

    def _set_cell_bg(self, cell, color: RGBColor):
        """Set background color of a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), str(color))
        tcPr.append(shd)
